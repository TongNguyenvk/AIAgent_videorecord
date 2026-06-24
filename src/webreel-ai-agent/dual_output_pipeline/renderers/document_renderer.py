"""
Document Renderer - Tao file DOCX tu plan va screenshots
Chuẩn văn bản học thuật/doanh nghiệp
"""
import logging
import re
from pathlib import Path
from typing import Dict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import base_renderer without relative import
import sys
sys.path.insert(0, str(Path(__file__).parent))
from base_renderer import Renderer

logger = logging.getLogger(__name__)

class DocumentRenderer(Renderer):
    """Renderer de tao file DOCX theo chuẩn học thuật/doanh nghiệp"""
    
    def render(self, plan: Dict, artifacts: Dict) -> str:
        """Tao file DOCX tu plan va screenshots"""
        logger.info("  [DocumentRenderer] Bat dau render DOCX...")
        
        doc = Document()
        self._setup_styles(doc)
        
        # Them tieu de chinh (Title/Heading 0)
        title = plan.get('title', 'Tutorial')
        self._add_title(doc, title)
        
        # Them mo ta (Introduction)
        intro_text = 'Tài liệu hướng dẫn này được tạo tự động bởi WebReel AI Agent. Vui lòng làm theo từng bước để hoàn thành tác vụ.'
        self._add_introduction(doc, intro_text)
        
        doc.add_paragraph()  # Khoang trang
        
        # Them tung buoc
        steps = plan.get('steps', [])
        screenshots = artifacts.get('screenshots', [])
        
        for i, step in enumerate(steps):
            # Lay narration va clean up
            narration = step.get('narration', f'Buoc {i+1}')
            
            # Remove [NARRATION:X] prefix neu co
            narration = re.sub(r'\[NARRATION:\d+\]\s*', '', narration)
            
            # Lay screenshot path
            screenshot_path = screenshots[i] if i < len(screenshots) else None
            
            # Them buoc voi format chuan
            self._add_step(doc, i+1, narration, screenshot_path)
        
        # Them footer
        self._add_footer(doc)
        
        # Luu file
        output_path = self.output_dir / f"{plan.get('name', 'tutorial')}.docx"
        doc.save(output_path)
        
        logger.info(f"  [DocumentRenderer] Da luu DOCX: {output_path}")
        return str(output_path)
    
    def _setup_styles(self, doc):
        """
        Thiet lap style mac dinh cho document
        Style Mac dinh (Normal Style): Times New Roman, 13 pt
        """
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(13)
        
        logger.info("  [DocumentRenderer] Setup styles: Times New Roman 13pt")
    
    def _add_title(self, doc, title: str):
        """
        Them tieu de chinh (Title/Heading 0)
        - Font size 16 pt
        - In dam (bold)
        - Canh giua (CENTER)
        - Mau den (khong dung mau xanh mac dinh)
        """
        heading = doc.add_heading(level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = heading.add_run(title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Mau den
        
        # Xoa default style cua heading
        heading.style.font.color.rgb = RGBColor(0, 0, 0)
    
    def _add_introduction(self, doc, intro_text: str):
        """
        Them phan mo ta
        - Canh deu hai ben (JUSTIFY)
        - Space after 12 pt
        - Italic
        """
        intro = doc.add_paragraph()
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        intro.paragraph_format.space_after = Pt(12)
        
        run = intro.add_run(intro_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.font.italic = True
    
    def _add_step(self, doc, step_number: int, narration: str, screenshot_path: str = None):
        """
        Them mot buoc voi format chuan
        
        Args:
            step_number: So thu tu buoc (1, 2, 3...)
            narration: Loi thoai/huong dan chi tiet
            screenshot_path: Duong dan anh chup man hinh (optional)
        
        Format:
        - Tieu de buoc: "Bước [Index]:" (chi tieu de, khong kem noi dung)
        - In dam, size 14 pt, mau den
        - Noi dung: Xuong hang moi, canh deu hai ben (JUSTIFY), space_after 12 pt
        - Anh: width=6.0 inches, canh giua
        """
        # Clean narration: Remove "Bước X" neu co o dau
        clean_narration = re.sub(r'^Bước\s+\d+[:\s]*', '', narration, flags=re.IGNORECASE).strip()
        
        # Neu narration qua ngan (< 10 chars), dung narration goc
        if len(clean_narration) < 10:
            clean_narration = narration.strip()
        
        # Tieu de buoc (Heading 1)
        # Format: Chi hien thi "Bước X:" (khong kem noi dung)
        step_heading = doc.add_heading(level=1)
        step_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        run = step_heading.add_run(f'Bước {step_number}:')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        step_heading.style.font.color.rgb = RGBColor(0, 0, 0)
        
        # Noi dung chi tiet (Paragraph text)
        # Hien thi full narration o paragraph ben duoi
        desc_para = doc.add_paragraph()
        desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        desc_para.paragraph_format.space_after = Pt(12)
        
        run = desc_para.add_run(clean_narration)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        
        # Them anh neu co
        if screenshot_path and Path(screenshot_path).exists():
            try:
                # Tao paragraph moi cho anh
                pic_para = doc.add_paragraph()
                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Canh giua
                
                # Them anh voi width=6.0 inches (ep cung)
                run = pic_para.add_run()
                run.add_picture(screenshot_path, width=Inches(6.0))
                
                # Them caption (optional, canh giua)
                caption = doc.add_paragraph(f'Hình {step_number}: Minh họa bước {step_number}')
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(12)
                
                run = caption.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)  # Mau xam nhe
                
            except Exception as e:
                logger.error(f"  [DocumentRenderer] Loi chen anh: {e}")
        
        # Them khoang trang giua cac buoc
        doc.add_paragraph()
    
    def _add_footer(self, doc):
        """
        Them footer cho document
        - Canh giua
        - Font size 11 pt
        - Mau xam
        """
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = 'WebReel AI Agent - Tài liệu hướng dẫn'
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in footer_para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(128, 128, 128)
    
    @property
    def output_format(self) -> str:
        return "docx"


